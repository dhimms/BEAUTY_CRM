# -*- coding: utf-8 -*-
"""
Generator Laporan PKL BEAUTY_CRM
Menghasilkan:
1. docs/Laporan_PKL_BEAUTY_CRM.docx
2. docs/Laporan_PKL_BEAUTY_CRM.md

BAB I dan BAB II mengikuti teks persis dari laporan referensi PT Microdata Indonesia,
disesuaikan dengan tanggal PKL 6 Juli 2026 - 6 September 2026 dan nama proyek BEAUTY_CRM.
BAB III sangat detail dan komprehensif memuat seluruh fitur teknis tim.
"""

import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def generate_report():
    doc = Document()

    # Set Margins (Standard Indonesian Thesis/Internship: Top 4cm, Left 4cm, Bottom 3cm, Right 3cm)
    for section in doc.sections:
        section.top_margin = Inches(1.57)     # ~4 cm
        section.left_margin = Inches(1.57)    # ~4 cm
        section.bottom_margin = Inches(1.18)  # ~3 cm
        section.right_margin = Inches(1.18)   # ~3 cm
        section.page_width = Inches(8.27)     # A4
        section.page_height = Inches(11.69)   # A4

    # Base normal style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_title(text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        return p

    def add_p(text, bold_prefix=None, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.5
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.name = 'Times New Roman'
            r_pre.font.size = Pt(12)
        r = p.add_run(text)
        r.italic = italic
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        return p

    def add_placeholder_box(caption, desc="[Placeholder Tangkapan Layar / Diagram]"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"┌──────────────────────────────────────────────────────────┐\n│ {desc} │\n└──────────────────────────────────────────────────────────┘")
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)

        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        r_cap = p_cap.add_run(caption)
        r_cap.italic = True
        r_cap.bold = True
        r_cap.font.name = 'Times New Roman'
        r_cap.font.size = Pt(11)

    def set_cell_shading(cell, color_hex="F2F4F7"):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_table_borders(table):
        tblPr = table._tbl.tblPr
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="333333"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="333333"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    def add_custom_table(headers, rows, caption=None, col_widths=None):
        if caption:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_cap.paragraph_format.space_before = Pt(8)
            p_cap.paragraph_format.space_after = Pt(3)
            r_cap = p_cap.add_run(caption)
            r_cap.bold = True
            r_cap.font.name = 'Times New Roman'
            r_cap.font.size = Pt(11)

        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)

        # Header
        hdr_cells = table.rows[0].cells
        for idx, header_text in enumerate(headers):
            cell = hdr_cells[idx]
            cell.text = header_text
            set_cell_shading(cell, "E5E7EB")
            set_cell_margins(cell, 120, 120, 150, 150)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.name = 'Times New Roman'
            p.runs[0].font.size = Pt(10.5)

        # Rows
        for r_idx, row_data in enumerate(rows):
            row_cells = table.rows[r_idx + 1].cells
            for c_idx, val in enumerate(row_data):
                cell = row_cells[c_idx]
                cell.text = str(val)
                set_cell_margins(cell, 100, 100, 150, 150)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 or (len(str(val)) < 12 and c_idx in [1, 2]) else WD_ALIGN_PARAGRAPH.LEFT
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(10)

        if col_widths:
            for row in table.rows:
                for idx, width in enumerate(col_widths):
                    row.cells[idx].width = width

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    print("Building Document Structure...")

    # ─────────────────────────────────────────────────────────────
    # HALAMAN SAMPUL / COVER
    # ─────────────────────────────────────────────────────────────
    add_title("LAPORAN PRAKTIK KERJA LAPANGAN", size=14, bold=True, space_before=10, space_after=18)
    add_title("RANCANG BANGUN SISTEM CUSTOMER RELATIONSHIP MANAGEMENT (BEAUTY_CRM) BERBASIS WEB DENGAN ARSITEKTUR MULTI-ROLE DAN LAYANAN CONTAINERIZED PADA PT MICRODATA INDONESIA", size=13, bold=True, space_after=36)
    
    add_title("[Placeholder: Logo Universitas Teknokrat Indonesia]", size=11, bold=False, space_after=36)
    
    add_title("Disusun Oleh:", size=12, bold=True, space_after=6)
    add_title("1. I NYOMAN VIVEKA\t(NPM: 23312003)\n2. ALVIN SAPUTRA\t(NPM: 23312138)\n3. DIMAS APRIANTO\t(NPM: 23312136)", size=12, bold=False, space_after=40)
    
    add_title("PROGRAM STUDI S1 INFORMATIKA\nFAKULTAS TEKNIK DAN ILMU KOMPUTER\nUNIVERSITAS TEKNOKRAT INDONESIA\nBANDAR LAMPUNG\n2026", size=13, bold=True, space_before=30)
    
    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # LEMBAR PERSETUJUAN
    # ─────────────────────────────────────────────────────────────
    add_heading_1("LEMBAR PERSETUJUAN")
    add_title("LAPORAN PRAKTIK KERJA LAPANGAN", size=12, bold=True, space_after=12)
    add_title("RANCANG BANGUN SISTEM CUSTOMER RELATIONSHIP MANAGEMENT (BEAUTY_CRM) BERBASIS WEB DENGAN ARSITEKTUR MULTI-ROLE DAN LAYANAN CONTAINERIZED PADA PT MICRODATA INDONESIA", size=11, bold=True, space_after=18)

    add_p("Disusun Oleh:", bold_prefix="", space_after=4)
    add_p("1. I Nyoman Viveka\tNPM: 23312003\tProgram Studi: S1 Informatika\n"
          "2. Alvin Saputra\tNPM: 23312138\tProgram Studi: S1 Informatika\n"
          "3. Dimas Aprianto\tNPM: 23312136\tProgram Studi: S1 Informatika", space_after=18)

    add_p("Laporan Praktik Kerja Lapangan ini telah diperiksa dan disetujui untuk diujikan pada Seminar Praktik Kerja Lapangan Program Studi S1 Informatika Fakultas Teknik dan Ilmu Komputer Universitas Teknokrat Indonesia.", space_after=24)

    add_p("Bandar Lampung, [ISI TANGGAL ACC, misal: 15 September 2026]", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

    tbl_persetujuan = doc.add_table(rows=3, cols=2)
    tbl_persetujuan.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_p1 = tbl_persetujuan.rows[0].cells[0]
    cell_p2 = tbl_persetujuan.rows[0].cells[1]
    cell_p1.text = "Menyetujui,\nPembimbing Lapangan\nPT Microdata Indonesia\n\n\n\nSigit Wasis Subekti\nProgrammer"
    cell_p2.text = "Menyetujui,\nDosen Pembimbing PKL\nUniversitas Teknokrat Indonesia\n\n\n\n[Nama Dosen Pembimbing, S.Kom., M.Kom.]\nNIDN. [NIDN Dosen Pembimbing]"
    for row in tbl_persetujuan.rows:
        for c in row.cells:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraphs[0].paragraph_format.line_spacing = 1.15
            for r in c.paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # LEMBAR PENGESAHAN
    # ─────────────────────────────────────────────────────────────
    add_heading_1("LEMBAR PENGESAHAN")
    add_title("LAPORAN PRAKTIK KERJA LAPANGAN", size=12, bold=True, space_after=12)
    add_title("RANCANG BANGUN SISTEM CUSTOMER RELATIONSHIP MANAGEMENT (BEAUTY_CRM) BERBASIS WEB DENGAN ARSITEKTUR MULTI-ROLE DAN LAYANAN CONTAINERIZED PADA PT MICRODATA INDONESIA", size=11, bold=True, space_after=18)

    add_p("Disusun Oleh:", space_after=4)
    add_p("1. I Nyoman Viveka\t(NPM: 23312003)\n2. Alvin Saputra\t(NPM: 23312138)\n3. Dimas Aprianto\t(NPM: 23312136)", space_after=18)

    add_p("Telah dipertahankan di hadapan Dewan Penguji Seminar Praktik Kerja Lapangan Program Studi S1 Informatika Fakultas Teknik dan Ilmu Komputer Universitas Teknokrat Indonesia pada tanggal [ISI TANGGAL SEMINAR PKL] dan dinyatakan LULUS.", space_after=24)

    tbl_pengesahan = doc.add_table(rows=4, cols=2)
    tbl_pengesahan.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pengesahan.rows[0].cells[0].text = "Dewan Penguji:\nPenguji I / Dosen Pembimbing\n\n\n\n[Nama Dosen Pembimbing, S.Kom., M.Kom.]\nNIDN. [NIDN Dosen]"
    tbl_pengesahan.rows[0].cells[1].text = "Pembimbing Lapangan\nPT Microdata Indonesia\n\n\n\nSigit Wasis Subekti\nProgrammer"
    
    tbl_pengesahan.rows[2].cells[0].text = "Mengetahui,\nKetua Program Studi S1 Informatika\n\n\n\nDidi Supriyadi, S.Kom., M.Kom.\nNIK. 022 10 05 08"
    tbl_pengesahan.rows[2].cells[1].text = "Mengesahkan,\nDekan Fakultas Teknik dan Ilmu Komputer\n\n\n\nDr. H. Mahathir Muhammad, S.E., M.M.\nNIK. 021 04 02 01"

    for row in tbl_pengesahan.rows:
        for c in row.cells:
            if c.paragraphs[0].text:
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                c.paragraphs[0].paragraph_format.line_spacing = 1.15
                for r in c.paragraphs[0].runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # RINGKASAN PELAKSANAAN PKL
    # ─────────────────────────────────────────────────────────────
    add_heading_1("RINGKASAN PELAKSANAAN PRAKTIK KERJA LAPANGAN")
    
    add_p("Praktik Kerja Lapangan (PKL) dilaksanakan di PT Microdata Indonesia, sebuah perusahaan terpercaya yang menyediakan layanan Software Developer, Digital Marketing, IT Consultant, dan Multimedia di Bandar Lampung. Kegiatan PKL berlangsung selama dua bulan, terhitung mulai tanggal 6 Juli 2026 sampai dengan 6 September 2026. Tim mahasiswa yang beranggotakan I Nyoman Viveka (NPM 23312003), Alvin Saputra (NPM 23312138), dan Dimas Aprianto (NPM 23312136) ditugaskan pada Divisi Web Development di bawah bimbingan Bapak Sigit Wasis Subekti selaku Programmer untuk mengembangkan sistem aplikasi berbasis web bertajuk BEAUTY_CRM (Beauty Customer Relationship Management).")
    
    add_p("BEAUTY_CRM merupakan solusi perangkat lunak CRM terintegrasi yang dirancang khusus untuk memenuhi kebutuhan operasional dan tata kelola hubungan pelanggan pada industri kecantikan, seperti klinik estetika medik, salon kecantikan, beauty studio, dan spa. Sistem ini dibangun dengan arsitektur Model-View-Controller (MVC) yang diperkuat dengan Service Layer Pattern pada framework Laravel 12, didukung oleh Tailwind CSS 4, Alpine.js 3, dan Vite 7. Sistem menerapkan kontrol akses berbasis peran (Role-Based Access Control) menggunakan Spatie Permission 6.x dengan empat tingkatan hak akses: Administrator, Sales, Customer Service (CS), dan Manager.")

    add_p("Dalam pelaksanaannya, pembagian tugas pengembangan sistem dibagi secara profesional berdasarkan kapabilitas masing-masing anggota:")
    add_p("1. I Nyoman Viveka berfokus pada pengembangan Modul Sales, yang mencakup manajemen prospek (leads management), kualifikasi lead, konversi lead menjadi deal transaksi, perancangan papan Kanban interaktif dengan teknologi SortableJS untuk pemindahan tahapan pipeline secara drag-and-drop, penutupan transaksi (deal won/lost), pencatatan log aktivitas dan penjadwalan tindak lanjut (activity & follow-up tracking), serta penyiaran pesan penjualan (sales deal blast).")
    add_p("2. Alvin Saputra berfokus pada pengembangan Modul Administrator dan System Core, yang bertanggung jawab atas manajemen akun staf dan penetapan target bulanan (user & target management), pengelolaan master data sumber prospek dan alasan kegagalan transaksi, fitur penataan ulang urutan pipeline (pipeline stage reordering), sistem impor dan ekspor data prospek massal berbasis Excel (Maatwebsite Excel) dengan perbaikan penanganan data numerik, sistem pencatatan rekam jejak otomatis (automated polymorphic audit logging via Observer Pattern), integrasi penyimpanan cloud object storage MinIO (S3 compatible), serta orkestrasi lingkungan kontainer Docker Compose.")
    add_p("3. Dimas Aprianto berfokus pada pengembangan Modul Manager dan Customer Service, yang mencakup perancangan Business Intelligence & Reports Center (ReportService) untuk menyajikan visualisasi performa penjualan, tren pendapatan 12 bulan, rasio konversi (win rate), durasi closing rata-rata, mesin peramalan pendapatan masa depan berbasis nilai terbobot (Weighted Pipeline Forecasting Engine), manajemen basis data pelanggan dan segmentasi belanja (min spend), sistem penjadwalan tindak lanjut dan pengingat keterlambatan (follow-up & overdue scheduler), serta mesin broadcast pesan multi-saluran terintegrasi via WhatsApp Gateway (Fonnte API) dan Email SMTP dengan teknologi Content-ID (CID) inline image embedding.")

    add_p("Hasil pelaksanaan PKL menunjukkan bahwa sistem BEAUTY_CRM berhasil dibangun, diuji, dan diimplementasikan secara komprehensif sesuai dengan kebutuhan industri kecantikan modern. Sistem ini mampu mengotomatisasi konversi prospek menjadi pelanggan loyal, menyajikan transparansi kinerja tim bagi level manajerial, serta memberikan keamanan data yang tinggi melalui arsitektur kontainerisasi dan audit trail.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # KATA PENGANTAR
    # ─────────────────────────────────────────────────────────────
    add_heading_1("KATA PENGANTAR")
    
    add_p("Puji dan syukur penulis panjatkan ke hadirat Allah SWT, Tuhan Yang Maha Esa, karena atas berkat, rahmat, dan karunia-Nya, penyusunan Laporan Praktik Kerja Lapangan (PKL) dengan judul \"Rancang Bangun Sistem Customer Relationship Management (BEAUTY_CRM) Berbasis Web dengan Arsitektur Multi-Role dan Layanan Containerized pada PT Microdata Indonesia\" dapat diselesaikan dengan baik dan tepat pada waktunya.")

    add_p("Laporan ini disusun sebagai salah satu syarat kelulusan mata kuliah Praktik Kerja Lapangan pada Program Studi S1 Informatika, Fakultas Teknik dan Ilmu Komputer, Universitas Teknokrat Indonesia. Selama pelaksanaan PKL yang dimulai pada tanggal 6 Juli 2026 hingga 6 September 2026 di PT Microdata Indonesia, penulis telah memperoleh banyak pengalaman berharga, pengetahuan praktis, dan wawasan industri mengenai pengembangan perangkat lunak berskala profesional.")

    add_p("Keberhasilan pelaksanaan PKL dan penyusunan laporan ini tidak lepas dari bimbingan, dorongan, arahan, dan bantuan dari berbagai pihak. Oleh karena itu, pada kesempatan ini penulis ingin menyampaikan rasa terima kasih dan apresiasi yang sebesar-besarnya kepada:")
    add_p("1. Bapak Dr. H. M. Nasrullah Yusuf, S.E., M.B.A., selaku Rektor Universitas Teknokrat Indonesia.\n"
          "2. Bapak Dr. H. Mahathir Muhammad, S.E., M.M., selaku Dekan Fakultas Teknik dan Ilmu Komputer Universitas Teknokrat Indonesia.\n"
          "3. Bapak Didi Supriyadi, S.Kom., M.Kom., selaku Ketua Program Studi S1 Informatika Universitas Teknokrat Indonesia.\n"
          "4. Bapak [Nama Dosen Pembimbing, S.Kom., M.Kom.], selaku Dosen Pembimbing PKL yang telah meluangkan waktu untuk memberikan bimbingan, arahan, dan masukan yang berharga dalam penyusunan laporan ini.\n"
          "5. Seluruh Bapak dan Ibu Dosen Program Studi S1 Informatika Universitas Teknokrat Indonesia yang telah membekali ilmu pengetahuan selama masa perkuliahan.\n"
          "6. Bapak Direktur dan Manajemen PT Microdata Indonesia yang telah memberikan izin dan kesempatan kepada penulis untuk melaksanakan PKL di perusahaan.\n"
          "7. Bapak Sigit Wasis Subekti, selaku Pembimbing Lapangan dan Programmer di PT Microdata Indonesia yang senantiasa memberikan bimbingan teknis, mentoring, serta transfer ilmu yang sangat bermanfaat selama proses pengembangan proyek.\n"
          "8. Seluruh rekan kerja dan staf PT Microdata Indonesia atas keramahan, suasana kerja yang kondusif, dan kerja sama yang baik selama masa PKL.\n"
          "9. Orang tua dan keluarga tercinta yang senantiasa memberikan doa tulus, motivasi tanpa henti, serta dukungan moral maupun material kepada penulis.\n"
          "10. Rekan-rekan mahasiswa Program Studi S1 Informatika Universitas Teknokrat Indonesia angkatan 2023 atas kebersamaan, bantuan, dan semangat perjuangan yang saling menguatkan.")

    add_p("Penulis menyadari bahwa penyusunan laporan ini masih jauh dari kesempurnaan. Oleh karena itu, kritik dan saran yang bersifat membangun sangat diharapkan guna penyempurnaan di masa yang akan datang. Semoga laporan ini dapat memberikan manfaat, wawasan, dan kontribusi positif bagi pembaca, institusi pendidikan, dan perkembangan ilmu rekayasa perangkat lunak.")

    add_p("Bandar Lampung, September 2026\n\n\nPenulis", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # DAFTAR ISI, DAFTAR TABEL, DAFTAR GAMBAR
    # ─────────────────────────────────────────────────────────────
    add_heading_1("DAFTAR ISI")
    add_p("HALAMAN JUDUL ......................................................................................................... i\n"
          "LEMBAR PERSETUJUAN ............................................................................................... ii\n"
          "LEMBAR PENGESAHAN ................................................................................................ iii\n"
          "RINGKASAN PELAKSANAAN PKL ............................................................................ iv\n"
          "KATA PENGANTAR ........................................................................................................ v\n"
          "DAFTAR ISI ....................................................................................................................... vi\n"
          "DAFTAR TABEL ............................................................................................................... viii\n"
          "DAFTAR GAMBAR .......................................................................................................... ix\n"
          "DAFTAR LAMPIRAN ...................................................................................................... xi\n\n"
          "BAB I PENDAHULUAN ................................................................................................. 1\n"
          "  1.1 Latar Belakang .......................................................................................................... 1\n"
          "  1.2 Tujuan ...................................................................................................................... 2\n"
          "  1.3 Manfaat .................................................................................................................... 3\n"
          "  1.4 Tempat Pelaksanaan ................................................................................................ 4\n"
          "  1.5 Jadwal Pelaksanaan ................................................................................................. 5\n\n"
          "BAB II TINJAUAN UMUM TEMPAT ......................................................................... 6\n"
          "  2.1 Profil Perusahaan ..................................................................................................... 6\n"
          "    2.1.1 Visi Perusahaan .................................................................................................. 7\n"
          "    2.1.2 Misi Perusahaan ................................................................................................. 7\n"
          "  2.2 Struktur Organisasi ................................................................................................... 7\n\n"
          "BAB III PELAKSANAAN PRAKTIK KERJA LAPANGAN ..................................... 10\n"
          "  3.1 Pelaksanaan PKL: I Nyoman Viveka, NPM 23312003 ............................................ 10\n"
          "    3.1.1 Bidang Kerja ...................................................................................................... 10\n"
          "    3.1.2 Pelaksanaan Kerja .............................................................................................. 11\n"
          "      3.1.2.1 Analisis & Perancangan Modul Penjualan (Sales Module) ........................ 11\n"
          "      3.1.2.2 Implementasi Siklus Hidup & Kualifikasi Prospek (Leads) ...................... 12\n"
          "      3.1.2.3 Konversi Prospek ke Transaksi (Lead-to-Deal Conversion) ...................... 14\n"
          "      3.1.2.4 Kanban Pipeline Interaktif & Drag-and-Drop SortableJS ......................... 16\n"
          "      3.1.2.5 Penutupan Transaksi (Deal Won/Lost) & Pembuatan Customer ................. 18\n"
          "      3.1.2.6 Pencatatan Interaksi & Pengingat (Activity & Follow-Up) ......................... 20\n"
          "      3.1.2.7 Penyiaran Pesan Penjualan (Sales Deal Blast Messaging) ........................ 22\n"
          "      3.1.2.8 Struktur Routing, Form Request, dan Model Modul Sales ....................... 23\n"
          "    3.1.3 Kendala yang Dihadapi ....................................................................................... 26\n"
          "    3.1.4 Solusi Penyelesaian Kendala .............................................................................. 27\n\n"
          "  3.2 Pelaksanaan PKL: Alvin Saputra, NPM 23312138 ................................................... 29\n"
          "    3.2.1 Bidang Kerja ...................................................................................................... 29\n"
          "    3.2.2 Pelaksanaan Kerja .............................................................................................. 30\n"
          "      3.2.2.1 Analisis & Perancangan Kontrol Akses Berbasis Peran (RBAC) ............... 30\n"
          "      3.2.2.2 Manajemen Pengguna & Penetapan Target (User & Target Management) . 32\n"
          "      3.2.2.3 Pengelolaan Data Master & Reordering Tahapan Pipeline ........................ 34\n"
          "      3.2.2.4 Impor dan Ekspor Data Massal Berbasis Excel (Maatwebsite Excel) ......... 36\n"
          "      3.2.2.5 Sistem Rekam Jejak Audit Otomatis (Polymorphic Audit Observer) ......... 38\n"
          "      3.2.2.6 Integrasi Cloud Object Storage MinIO (S3 Compatible Storage) ............... 40\n"
          "      3.2.2.7 Orkestrasi Kontainerisasi Sistem dengan Docker Compose ....................... 42\n"
          "      3.2.2.8 Struktur Routing, Form Request, dan Model Modul Administrator ........... 45\n"
          "    3.2.3 Kendala yang Dihadapi ....................................................................................... 48\n"
          "    3.2.4 Solusi Penyelesaian Kendala .............................................................................. 49\n\n"
          "  3.3 Pelaksanaan PKL: Dimas Aprianto, NPM 23312136 ................................................ 51\n"
          "    3.3.1 Bidang Kerja ...................................................................................................... 51\n"
          "    3.3.2 Pelaksanaan Kerja .............................................................................................. 52\n"
          "      3.3.2.1 Analisis Business Intelligence & Customer Service Workflow ................. 52\n"
          "      3.3.2.2 Perancangan Arsitektur Report Center (ReportService) ............................ 54\n"
          "      3.3.2.3 Metrik Performa Sales, Rasio Konversi, dan Tren Pendapatan ................... 56\n"
          "      3.3.2.4 Mesin Peramalan Pendapatan & Member (Weighted Pipeline Forecast) ..... 58\n"
          "      3.3.2.5 Basis Data Pelanggan, Profiling, dan Segmentasi Nilai Belanja ................ 60\n"
          "      3.3.2.6 Sistem Manajemen Jadwal Tindak Lanjut & Monitoring Overdue .............. 62\n"
          "      3.3.2.7 Integrasi Multi-Channel Blast (WhatsApp Fonnte API & Email CID) ....... 64\n"
          "      3.3.2.8 Struktur Routing, Form Request, dan Model Modul Manager & CS ......... 66\n"
          "    3.3.3 Kendala yang Dihadapi ....................................................................................... 69\n"
          "    3.3.4 Solusi Penyelesaian Kendala .............................................................................. 70\n\n"
          "BAB IV PENUTUP ......................................................................................................... 72\n"
          "  4.1 Kesimpulan .............................................................................................................. 72\n"
          "  4.2 Saran ........................................................................................................................ 74\n\n"
          "DAFTAR PUSTAKA ....................................................................................................... 76\n"
          "LAMPIRAN ...................................................................................................................... 79")

    doc.add_page_break()

    add_heading_1("DAFTAR TABEL")
    add_p("Table 1.1 Jadwal Pelaksanaan ................................................................................................. 5\n"
          "Tabel 3.1 Matriks Hak Akses Peran (Role & Permission Matrix) BEAUTY_CRM ................. 31\n"
          "Tabel 3.2 Daftar Endpoint Rute Modul Sales ........................................................................... 24\n"
          "Tabel 3.3 Daftar Endpoint Rute Modul Administrator ............................................................. 46\n"
          "Tabel 3.4 Daftar Endpoint Rute Modul Manager & Customer Service .................................... 67\n"
          "Tabel 3.6 Konfigurasi Port & Layanan Docker Compose ....................................................... 43\n"
          "Tabel 3.7 Parameter Konfigurasi Environment Object Storage MinIO S3 ............................. 41\n"
          "Tabel 3.8 Parameter Probabilitas Tahapan Pipeline Deals ....................................................... 59")

    doc.add_page_break()

    add_heading_1("DAFTAR GAMBAR")
    add_p("Gambar 1.1 Gambar Map Lokasi PKL ................................................................................. 4\n"
          "Gambar 2.1 Logo Perusahan .................................................................................................... 6\n"
          "Gambar 2.2 Struktur Organisasi Perusahaan ............................................................................ 7\n"
          "Gambar 3.1 Arsitektur Aplikasi BEAUTY_CRM (MVC + Service Layer) ............................. 12\n"
          "Gambar 3.2 Alur Siklus Hidup Prospek (Lead Lifecycle Flowchart) ........................................ 13\n"
          "Gambar 3.3 Tampilan Antarmuka Daftar Prospek (Leads List View) ..................................... 14\n"
          "Gambar 3.4 Modal Kualifikasi dan Konversi Prospek Menjadi Transaksi Deal ...................... 16\n"
          "Gambar 3.5 Tampilan Papan Interaktif Kanban Pipeline Sales (SortableJS) ............................ 17\n"
          "Gambar 3.6 Modal Penutupan Transaksi (Close Deal as Won / Lost) ...................................... 19\n"
          "Gambar 3.7 Timeline Riwayat Aktivitas & Follow-Up Interaksi Penjualan ............................ 21\n"
          "Gambar 3.8 Form Siaran Pesan Penjualan (Sales Blast Modal) ............................................... 23\n"
          "Gambar 3.9 Antarmuka Manajemen Akun Pengguna & Pengaturan Target Performa .............. 33\n"
          "Gambar 3.10 Antarmuka Penataan Ulang Tahapan Pipeline (Drag & Drop Reorder) ................ 35\n"
          "Gambar 3.11 Antarmuka Fitur Impor Data Prospek Massal via Excel ...................................... 37\n"
          "Gambar 3.12 Tampilan Tabel Riwayat Rekam Jejak Audit Sistem (Audit Trail) ..................... 39\n"
          "Gambar 3.13 Dashboard Manajemen Penyimpanan Objek MinIO Console ............................ 41\n"
          "Gambar 3.14 Diagram Arsitektur Kontainerisasi Layanan Docker Compose .......................... 44\n"
          "Gambar 3.15 Diagram Entity Relationship (ERD) Sistem BEAUTY_CRM ............................ 32\n"
          "Gambar 3.16 Tampilan Dashboard Utama Eksekutif Manager (BI Dashboard) ...................... 53\n"
          "Gambar 3.17 Laporan Analisis Performa Tim Sales & Conversion Win Rate ........................... 57\n"
          "Gambar 3.18 Diagram Analisis Tren Pendapatan 12 Bulan Historis ........................................ 58\n"
          "Gambar 3.19 Tampilan Antarmuka Mesin Peramalan (Weighted Forecasting Board) .............. 60\n"
          "Gambar 3.20 Antarmuka Manajemen Data Pelanggan & Segmentasi Belanja ......................... 61\n"
          "Gambar 3.21 Dashboard Jadwal Tindak Lanjut & Monitoring Keterlambatan CS ................... 63\n"
          "Gambar 3.22 Antarmuka Siaran Pesan Multi-Saluran CS (WhatsApp & Email Blast) ............ 65")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # BAB I PENDAHULUAN (PERSIS TEMPLATE REFERENSI DENGAN PENYESUAIAN RESMI)
    # ─────────────────────────────────────────────────────────────
    add_heading_1("BAB I\nPENDAHULUAN")
    
    add_heading_2("1.1 Latar Belakang")
    add_p("Praktik Kerja Lapangan (PKL) adalah komponen krusial dalam kurikulum pendidikan tinggi yang bertujuan memberikan pengalaman nyata kepada mahasiswa di lingkungan kerja. Melalui program PKL, mahasiswa mendapatkan peluang untuk menerapkan teori, konsep, dan pengetahuan yang telah dipelajari selama kuliah ke dalam praktik secara langsung di tempat kerja yang nyata. Oleh karena itu, terdapat hubungan antara aspek akademis yang diperoleh di perguruan tinggi dengan kebutuhan serta tantangan yang dihadapi oleh dunia industri dan lembaga profesional.")

    add_p("Selain berfungsi sebagai wadah penerapan ilmu pengetahuan, PKL juga memainkan peran penting dalam meningkatkan kemampuan mahasiswa, baik dalam hal keterampilan teknis (hard skills) maupun keterampilan nonteknis (soft skills) (Purbo, Y. S., Utomo, F. S., & Purwati, 2023). seperti kemampuan berkomunikasi, kerja sama, disiplin, tanggung jawab, dan keterampilan dalam menyelesaikan masalah yang ada di tempat kerja. Pengalaman yang didapat selama PKL diharapkan bisa menjadi modal bagi mahasiswa dalam mempersiapkan diri menghadapi kompetisi di dunia kerja setelah menyelesaikan pendidikan.")

    add_p("Sejalan dengan kemajuan teknologi informasi yang kian pesat, permintaan akan tenaga kerja yang memiliki keterampilan dan pengalaman kerja semakin bertambah (Ariawan, M. D., Triayudi, A., 2020). Maka dari itu, universitas bekerja sama dengan beragam perusahaan dan lembaga untuk memberikan peluang kepada mahasiswa mendapatkan pengalaman kerja yang sesuai dengan disiplin ilmunya. Kegiatan PKL merupakan salah satu wujud kerjasama tersebut untuk meningkatkan kualitas lulusan yang siap memenuhi kebutuhan industri.")

    add_p("Kegiatan PKL yang dilakukan penulis bertujuan untuk menerapkan pengetahuan yang telah diperoleh selama kuliah melalui partisipasi langsung dalam aktivitas kerja profesional. Pada kesempatan ini, penulis melakukan PKL di PT Microdata Indonesia untuk membangun sistem Customer Relationship Management (BEAUTY_CRM) berbasis web. Penguatan fitur interaksi prospek, otomatisasi pipeline transaksi, penyesuaian layanan purna jual, serta kemudahan sistem sangat menentukan keputusan pengguna dalam bertransaksi secara digital (Mal, L. H., & Mertayasa, 2018). Melalui aktivitas ini, penulis diharapkan bisa mendapatkan pengalaman kerja, meningkatkan keterampilan teknis serta profesional, serta memahami proses kerja dan budaya organisasi yang diterapkan di lingkungan perusahaan.")

    add_heading_2("1.2 Tujuan")
    add_p("Tujuan dari program PKL yang diikuti oleh penulis, antara lain:")
    add_p("1. Mengaplikasikan pengetahuan dan keterampilan yang diperoleh selama perkuliahan dalam pengembangan sistem informasi di PT Microdata Indonesia.\n"
          "2. Memperoleh pemahaman mengenai proses kerja di bidang teknologi informasi, khususnya pada pengembangan sistem Customer Relationship Management (BEAUTY_CRM).\n"
          "3. Mengasah kemampuan profesional melalui peningkatan keterampilan teknis, serta mengembangkan soft skills seperti komunikasi, kolaborasi tim, kedisiplinan, dan kemampuan pemecahan masalah.\n"
          "4. Menambah wawasan tentang pengembangan sistem manajemen pelanggan secara online dan real time baik dari tampilan web maupun layanan background.")

    add_heading_2("1.3 Manfaat")
    add_p("Pelaksanaan PKL memberikan berbagai manfaat, baik bagi mahasiswa, Universitas Teknokrat Indonesia, maupun bagi instansi tempat PKL dilaksanakan. Manfaat tersebut antara lain:")
    add_p("1. Manfaat bagi Mahasiswa:\n"
          "   a) Menambah pengalaman dalam membuat aplikasi, khususnya menggunakan Laravel, Tailwind CSS, Alpine.js, Spatie Permission, Docker, dan MinIO S3, belajar membuat sistem yang saling terhubung agar data bisa diakses dan dikelola secara otomatis.\n"
          "   b) Mampu memahami alur pengelolaan prospek (leads), negosiasi transaksi (deals pipeline), dan layanan purna jual pelanggan (customer service).\n"
          "   c) Melatih kemampuan menganalisis kebutuhan pengguna, lalu merancang dan membangun solusi yang tepat untuk digunakan.\n"
          "   d) Belajar kerja sama tim, berbagi tugas dan tanggung jawab dengan teman kelompok agar proyek bisa selesai dengan baik.\n"
          "   e) Dapat menambah pengalaman baru antara dunia perkuliahan dengan dunia kerja agar tahu kebutuhan apa yang harus disiapkan setelah lulus nanti.")
    add_p("2. Manfaat bagi Universitas:\n"
          "   a) Membina hubungan kerja sama yang baik dengan dunia usaha dan instansi, sehingga membuka peluang bagi mahasiswa lain untuk melaksanakan PKL di masa mendatang.\n"
          "   b) Menjadi sarana evaluasi terhadap kesesuaian kurikulum dengan kebutuhan dunia kerja, khususnya di bidang informatika.\n"
          "   c) Meningkatkan nama baik universitas melalui kontribusi nyata mahasiswa dalam dunia industri.")
    add_p("3. Manfaat Bagi Instansi:\n"
          "   a) Menjalin kerja sama yang lebih erat dengan perguruan tinggi di bidang teknologi informasi.\n"
          "   b) Memperoleh ide, inovasi, dan perspektif baru dari mahasiswa yang sedang belajar serta mengikuti perkembangan teknologi terkini.")

    add_heading_2("1.4 Tempat Pelaksanaan")
    add_p("Pelaksanaan PKL dilaksanakan di PT. Microdata Indonesia yang beralamat di Jl. Endro Suratmin No. 52D, Way Dadi, Kecamatan Sukarame, Kota Bandar Lampung. PT. Microdata Indonesia merupakan perusahaan yang bergerak di bidang teknologi informasi dan pengembangan perangkat lunak yang menyediakan berbagai layanan digital, seperti pengembangan aplikasi web, aplikasi mobile, sistem informasi, serta solusi teknologi untuk kebutuhan bisnis dan instansi. Dokumentasi lokasi pelaksanaan PKL di PT. Microdata Indonesia dapat dilihat pada Gambar 1.1.")

    add_placeholder_box("Gambar 1.1 Gambar Map Lokasi PKL", "[Peta Lokasi Kantor PT Microdata Indonesia - Jl. Endro Suratmin No. 52D]")

    add_heading_2("1.5 Jadwal Pelaksanaan")
    add_p("Kegiatan Program Magang ini dilaksanakan selama 2 bulan dimulai dari 6 Juli – 6 September 2026. Jadwal kegiatan selama Magang berlangsung setiap hari kerja, mulai dari hari Senin hingga Jumat pada pukul 09.00 sampai 17.00 WIB dapat dilihat di Table 1.1.")

    add_custom_table(
        headers=["Hari", "Jam Masuk", "Istirahat", "Masuk", "Pulang"],
        rows=[
            ["Senin", "09.00 WIB", "12.00 WIB", "13.00 WIB", "17.00 WIB"],
            ["Selasa", "09.00 WIB", "12.00 WIB", "13.00 WIB", "17.00 WIB"],
            ["Rabu", "09.00 WIB", "12.00 WIB", "13.00 WIB", "17.00 WIB"],
            ["Kamis", "09.00 WIB", "12.00 WIB", "13.00 WIB", "17.00 WIB"],
            ["Jumat", "09.00 WIB", "11.45 WIB", "13.30 WIB", "17.00 WIB"]
        ],
        caption="Table 1.1 Jadwal Pelaksanaan"
    )

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # BAB II TINJAUAN UMUM TEMPAT (PERSIS 100% TEMPLATE REFERENSI)
    # ─────────────────────────────────────────────────────────────
    add_heading_1("BAB II\nTINJAUAN UMUM TEMPAT")
    
    add_heading_2("2.1 Profil Perusahaan")
    add_p("PT Microdata Indonesia adalah perusahaan terpercaya yang menyediakan layanan Software Developer, Digital Marketing, IT Consultant, dan Multimedia yang berkembang di Indonesia sejak tahun 2010. PT Microdata Indonesia berlandaskan pada tujuan memberikan layanan berkualitas tinggi dan menempatkan kepuasan pelanggan sebagai prioritas utama, serta berkomitmen untuk selalu fokus dalam melayani dan menyediakan solusi terbaik sesuai kebutuhan pelanggan di bidang jasa telekomunikasi data. Sejalan dengan perkembangan teknologi yang semakin pesat di masa kini, “Teknologi Informasi” telah menjadi kebutuhan utama masyarakat secara umum. PT Microdata Indonesia berkomitmen untuk memberikan solusi layanan jasa yang bernilai tambah dengan menempatkan kualitas dan kepuasan pelanggan sebagai prioritas utama, agar perusahaan dapat meraih keberhasilan yang diharapkan. Tenaga muda yang dinamis dan berdedikasi tinggi mendukung PT Microdata Indonesia dalam memberikan pelayanan terbaik serta memiliki pengalaman luas di dunia teknologi informasi. Kehadiran tenaga muda yang penuh semangat ini adalah jaminan bahwa PT Microdata Indonesia akan terus berkembang dan mengikuti arus perkembangan zaman. PT. Microdata Indonesia menawarkan layanan berkualitas berkat pengalaman yang dimiliki stafnya. Berikut ini merupakan logo yang mencerminkan identitas PT Microdata Indonesia sebagai perusahaan yang bergerak di bidang teknologi informasi dan sistem digital. Logo ini dirancang sebagai representasi visual dari nama dan karakter perusahaan dalam konteks profesionalisme di dunia teknologi, sebagaimana dapat dilihat pada Gambar 2.1.")

    add_placeholder_box("Gambar 2.1 Logo Perusahan", "[Logo Resmi PT Microdata Indonesia]")

    add_heading_3("2.1.1 Visi Perusahaan")
    add_p("Menyediakan layanan jasa dan produk terbaik kepada pelanggan dengan prioritas utama pada kepuasan mereka, yang mendukung pertumbuhan PT. Sebagai layanan jasa Teknologi Informasi, Microdata Indonesia adalah pilihan utama.")

    add_heading_3("2.1.2 Misi Perusahaan")
    add_p("1. Menyediakan layanan unggulan yang didukung oleh berbagai inovasi dan solusi jaringan telekomunikasi guna meningkatkan kualitas pelayanan kepada pelanggan.\n"
          "2. Membangun dan memperluas infrastruktur jaringan di seluruh kota di Indonesia menggunakan teknologi terkini.\n"
          "3. Konsistensi dalam pengembangan Sumber Daya Manusia adalah langkah penting menuju keberhasilan.\n"
          "4. Senantiasa menciptakan inovasi produk baru dalam bidang teknologi informasi.")

    add_heading_2("2.2 Struktur Organisasi")
    add_p("Operasional perusahaan sangat bergantung pada keberadaan dan fungsi Struktur Organisasi. Struktur organisasi ini perlu dirancang agar dapat meningkatkan koordinasi dan kolaborasi antar bagian perusahaan, sehingga tanggung jawab dan mutu kerja di setiap bagian dapat lebih baik. Struktur organisasi PT Microdata Indonesia disusun sebagaimana ditampilkan pada Gambar 2.2.")

    add_placeholder_box("Gambar 2.2 Struktur Organisasi Perusahaan", "[Bagan Struktur Organisasi PT Microdata Indonesia]")

    add_p("Dalam struktur organisasi PT Microdata Indonesia, terdapat tugas dan fungsi yang bertujuan mendukung proses operasional perusahaan sebagai berikut:\n"
          "1. Direktur\n"
          "   a. Mengawasi dan memimpin jalannya perusahaan agar berjalan lancar sesuai kebijakan dan tujuan yang telah disepakati perusahaan.\n"
          "   b. Membangun kerjasama yang harmonis dengan perusahaan eksternal baik dari sektor swasta maupun pemerintah.\n"
          "   c. Mengembangkan serta mengesahkan kebijakan perusahaan.\n"
          "   d. Bertanggung jawab atas pertumbuhan perusahaan dan mengelola semua kegiatan operasionalnya.\n"
          "2. Manajer Operasional\n"
          "   a. Mengatur dan memperbaiki pertumbuhan serta kapabilitas perusahaan.\n"
          "   b. Mengawasi serta mengendalikan pengeluaran biaya agar sesuai dengan anggaran yang sudah ditentukan oleh perusahaan.\n"
          "   c. Melakukan analisis terhadap persoalan yang terjadi dalam sistem kegiatan perusahaan.\n"
          "   d. Mengoptimalkan sistem operasional serta kebijakan agar visi dan misi perusahaan dapat terealisasi dengan baik.\n"
          "   e. Memastikan proyek selesai tepat waktu sebagai tanggung jawab utama.\n"
          "3. Manajer Teknik\n"
          "   a. Mengelola bagian teknik untuk perencanaan dan pengembangan proyek.\n"
          "   b. Menyusun rencana tugas proyek.\n"
          "   c. Mengawasi jalannya survei serta pengumpulan data informasi.\n"
          "   d. Mengatur dan memperhatikan kegiatan teknik agar operasional perusahaan berjalan dengan lancar.\n"
          "4. Manajer Pemasaran\n"
          "   a. Membuat laporan kegiatan sebagai bentuk pertanggungjawaban kepada direktur perusahaan dan bertanggung jawab atas proses tersebut.\n"
          "   b. Meningkatkan jumlah penjualan baik secara langsung maupun melalui internet.\n"
          "   c. Mengelola media, acara, dan klien dalam rangka keperluan promosi demi meningkatkan volume penjualan.\n"
          "   d. Perusahaan fokus pada pengembangan jasa dan produk.\n"
          "5. Programmer\n"
          "   a. Membuat program yang diperlukan untuk memenuhi kebutuhan perusahaan.\n"
          "   b. Melaksanakan percobaan pada program aplikasi perangkat lunak guna memastikan bahwa output yang dihasilkan sesuai dengan yang diharapkan dan bahwa instruksi yang dieksekusi berjalan dengan benar.\n"
          "   c. Mengoreksi error dalam program dan melakukan pengecekan ulang untuk memastikan hasilnya memenuhi harapan.\n"
          "   d. Mengembangkan, memperbarui, serta menjaga program komputer dan perangkat lunak guna mendukung kegiatan operasional perusahaan.\n"
          "   e. Melaksanakan analisis sistem dan pemrograman guna pemeliharaan serta pengawasan pengguna perangkat.\n"
          "6. Database Administrator\n"
          "   a. Membangun dan mengubah database serta mengelola database yang telah tersedia.\n"
          "   b. Mengkoordinasikan pengembangan serta menetapkan proyek dan batasan pada database.\n"
          "   c. Menetapkan pengguna dan level akses untuk setiap segmen database.\n"
          "7. Analisis Sistem\n"
          "   a. Memilih perangkat lunak dan perangkat keras komputer yang digunakan dalam pengelolaan dan perubahan sistem.\n"
          "   b. Membangun diagram dan flowchart yang menggambarkan alur proses dari sistem yang ada.\n"
          "   c. Mengarsipkan solusi dan tantangan yang dihadapi oleh sistem yang aktif sebagai catatan untuk masa yang akan datang.\n"
          "   d. Mengumpulkan dan mencari data tentang sistem yang berjalan saat ini dan sistem yang akan dikembangkan di masa depan.\n"
          "   e. Perencanaan, studi, instalasi, troubleshooting, konfigurasi, pemeliharaan, serta peningkatan sistem operasi.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # BAB III PELAKSANAAN PRAKTIK KERJA LAPANGAN (DETAIL PROYEK KITA)
    # ─────────────────────────────────────────────────────────────
    add_heading_1("BAB III\nPELAKSANAAN PRAKTIK KERJA LAPANGAN")

    add_p("Bab ini menguraikan secara rinci dan komprehensif mengenai pelaksanaan Praktik Kerja Lapangan (PKL) yang dilaksanakan oleh tim mahasiswa di PT Microdata Indonesia dalam mengembangkan sistem BEAUTY_CRM. Untuk memberikan gambaran kerja yang spesifik, transparan, dan terukur, pembahasan dibagi menjadi tiga subbab utama yang masing-masing merepresentasikan kontribusi teknis, bidang kerja, implementasi modul, kendala yang dihadapi, serta solusi pemecahan masalah dari setiap anggota tim.")

    # ═════════════════════════════════════════════════════════════
    # SUBBAB 3.1 I NYOMAN VIVEKA
    # ═════════════════════════════════════════════════════════════
    add_heading_2("3.1 Pelaksanaan PKL: I NYOMAN VIVEKA, NPM 23312003")

    add_heading_3("3.1.1 Bidang Kerja")
    add_p("Pelaksanaan PKL yang dilakukan oleh I Nyoman Viveka (NPM 23312003) dilaksanakan di PT Microdata Indonesia pada Divisi Web Development di bawah bimbingan Bapak Sigit Wasis Subekti selaku Programmer. Penulis berfokus pada pengembangan Modul Penjualan (Sales Module) dan Pipeline Manajemen Transaksi. Penulis bertanggung jawab penuh terhadap perancangan dan implementasi siklus hidup prospek (lead lifecycle), otomasi kualifikasi calon pelanggan, mekanisme konversi prospek menjadi transaksi, visualisasi papan interaktif Kanban berbasis SortableJS, pencatatan log aktivitas interaksi negosiasi, hingga fitur broadcast pesan penawaran penjualan.")

    add_heading_3("3.1.2 Pelaksanaan Kerja")
    add_p("Pelaksanaan kerja yang dilakukan oleh I Nyoman Viveka dalam membangun Modul Sales BEAUTY_CRM diuraikan ke dalam beberapa tahapan fungsional sebagai berikut:")

    add_heading_3("3.1.2.1 Analisis & Perancangan Modul Penjualan (Sales Module)")
    add_p("Pada tahap awal, penulis melakukan analisis mendalam mengenai pola kerja staf penjualan di klinik kecantikan. Proses penjualan produk kecantikan dan paket perawatan estetika medik memiliki karakteristik berupa siklus pertimbangan yang membutuhkan konsultasi berulang. Oleh karena itu, arsitektur modul penjualan dirancang menggunakan pendekatan Lead-to-Deal Pipeline yang fleksibel dan interaktif.")

    add_placeholder_box("Gambar 3.1 Arsitektur Aplikasi BEAUTY_CRM (MVC + Service Layer)", "[Diagram Arsitektur MVC + Service Layer + Database + Storage]")

    add_heading_3("3.1.2.2 Implementasi Siklus Hidup & Kualifikasi Prospek (Leads Management)")
    add_p("Prospek calon pelanggan (leads) dapat masuk ke sistem melalui penugasan dari Administrator (via admin assignment / import Excel) atau diinputkan secara mandiri oleh staf Sales melalui formulir `sales.leads.create`. Prospek memiliki lima status utama yang bertransisi secara otomatis:")
    add_p("1. `new`: Prospek baru yang belum pernah dihubungi oleh staf sales.\n"
          "2. `contacted`: Status otomatis berubah ketika staf sales mencatat interaksi pertama (telepon/WhatsApp/email) pada lead terkait.\n"
          "3. `qualified`: Prospek dinilai memiliki minat nyata, anggaran yang sesuai, dan kebutuhan perawatan kecantikan yang valid.\n"
          "4. `converted`: Prospek yang telah berhasil dikonversi menjadi transaksi negosiasi aktif (Deal).\n"
          "5. `closed`: Prospek yang diarsipkan karena tidak berminat atau tidak dapat dihubungi.")

    add_placeholder_box("Gambar 3.2 Alur Siklus Hidup Prospek (Lead Lifecycle Flowchart)", "[Bagan Alir Status Lead: New -> Contacted -> Qualified -> Converted/Closed]")

    add_p("Pada controller `Sales\\LeadController`, penulis mengimplementasikan method `qualify()` yang menerima data dari `QualifyLeadRequest`. Kualifikasi dapat ditentukan menjadi `qualified`, `unqualified`, atau `not_fit` dengan catatan kualifikasi yang tersimpan di basis data.")

    add_placeholder_box("Gambar 3.3 Tampilan Antarmuka Daftar Prospek (Leads List View)", "[Screenshot Halaman Sales Leads Index dengan Filter Status & Pencarian]")

    add_heading_3("3.1.2.3 Konversi Prospek ke Transaksi (Lead-to-Deal Conversion)")
    add_p("Ketika prospek telah berada pada status `qualified`, staf sales dapat menekan tombol \"Convert to Deal\". Penulis membangun alur transaksi bisnis ini pada class `DealService::createFromLead()` yang dibungkus di dalam transaksi basis data (`DB::transaction`). Logika bisnis ini secara atomik membuat data entitas `Deal` baru pada tahapan pipeline pertama dan secara otomatis memperbarui status `Lead` menjadi `converted`:")

    add_p("```php\n"
          "public function createFromLead(Lead $lead, array $data): Deal {\n"
          "    return DB::transaction(function () use ($lead, $data) {\n"
          "        $firstStage = PipelineStage::ordered()->first();\n"
          "        $deal = Deal::create([\n"
          "            'lead_id'             => $lead->id,\n"
          "            'name'                => $data['name'],\n"
          "            'value'               => 0,\n"
          "            'pipeline_stage_id'   => $data['pipeline_stage_id'] ?? $firstStage->id,\n"
          "            'status'              => 'open',\n"
          "            'expected_close_date' => $data['expected_close_date'] ?? null,\n"
          "            'assigned_to'         => $lead->assigned_to,\n"
          "            'created_by'          => auth()->id(),\n"
          "        ]);\n"
          "        $lead->update(['status' => 'converted']);\n"
          "        return $deal;\n"
          "    });\n"
          "}\n"
          "```", italic=True)

    add_placeholder_box("Gambar 3.4 Modal Kualifikasi dan Konversi Prospek Menjadi Transaksi Deal", "[Screenshot Modal Dialog Convert to Deal]")

    add_heading_3("3.1.2.4 Kanban Pipeline Interaktif & Drag-and-Drop SortableJS")
    add_p("Salah satu kontribusi teknis paling signifikan dari penulis adalah implementasi visualisasi papan interaktif Kanban pada rute `/sales/pipeline` (`DealController::pipeline`). Papan ini memetakan seluruh transaksi aktif ke dalam kolom-kolom tahapan negosiasi (seperti Prospecting, Qualification, Treatment Proposal, Price Negotiation, dan Final Closing).")

    add_p("Penulis mengintegrasikan pustaka JavaScript `SortableJS` untuk memungkinkan staf sales menggeser kartu transaksi antar tahapan secara *drag-and-drop*. Ketika sebuah kartu transaksi dipindahkan ke kolom lain, event listener JavaScript mengirimkan request AJAX asinkronus ke endpoint `POST /sales/deals/{deal}/move-stage`:")

    add_p("```javascript\n"
          "new Sortable(stageColumnElement, {\n"
          "    group: 'deals-pipeline',\n"
          "    animation: 150,\n"
          "    onEnd: function (evt) {\n"
          "        let dealId = evt.item.dataset.dealId;\n"
          "        let targetStageId = evt.to.dataset.stageId;\n"
          "        fetch(`/sales/deals/${dealId}/move-stage`, {\n"
          "            method: 'POST',\n"
          "            headers: {\n"
          "                'Content-Type': 'application/json',\n"
          "                'X-CSRF-TOKEN': document.querySelector('meta[name=\"csrf-token\"]').content\n"
          "            },\n"
          "            body: JSON.stringify({ stage_id: targetStageId })\n"
          "        }).then(res => res.json()).then(data => {\n"
          "            updateStageBadgeAndTotals(data);\n"
          "        });\n"
          "    }\n"
          "});\n"
          "```", italic=True)

    add_placeholder_box("Gambar 3.5 Tampilan Papan Interaktif Kanban Pipeline Sales (SortableJS)", "[Screenshot Papan Kanban Pipeline Sales dengan Card Transaksi dan Badge Nilai]")

    add_heading_3("3.1.2.5 Penutupan Transaksi (Deal Won/Lost) & Pembuatan Customer Otomatis")
    add_p("Ketika proses negosiasi mencapai kesepakatan final, staf sales menutup transaksi melalui aksi \"Mark as Won\" atau \"Mark as Lost\" pada modal closing:")
    add_p("1. **Deal Closed as WON**: Sales memasukkan nama paket produk/treatment yang dibeli dan nilai transaksi final (`product_name` & `value`). Method `DealService::closeWon()` memperbarui status deal menjadi `won`, mencatat tanggal closing (`closed_at`), serta memeriksa apakah data `Customer` sudah pernah terdaftar dari `lead_id` terkait. Jika belum ada, sistem secara otomatis mengonversi data kontak prospek menjadi entitas data `Customer` baru berstatus `active` yang siap ditindaklanjuti oleh Customer Service. Selain itu, sistem menembakkan `DealWonNotification` ke seluruh Admin dan Manager.")
    add_p("2. **Deal Closed as LOST**: Jika kesepakatan gagal, sales wajib memilih alasan kegagalan (`lost_reason_id`) dari master data (misal: harga terlalu tinggi, memilih klinik lain, lokasi terlalu jauh) serta memberikan catatan alasan kekalahan (`lost_notes`). Data ini menjadi bahan evaluasi analitik bagi level manajemen.")

    add_placeholder_box("Gambar 3.6 Modal Penutupan Transaksi (Close Deal as Won / Lost)", "[Screenshot Dialog Form Mark as Won / Lost]")

    add_heading_3("3.1.2.6 Pencatatan Interaksi & Pengingat (Activity & Follow-Up)")
    add_p("Penulis mengimplementasikan modul pencatatan aktivitas interaksi pada `Sales\\ActivityController`. Model `Activity` bersifat polimorfik (`morphTo: activitable`) sehingga dapat dilekatkan pada entitas `Lead` maupun `Deal`. Staf sales dapat mencatat jenis interaksi (Call, WhatsApp, Email, Meeting/Konsultasi Tatap Muka, Catatan), durasi, hasil pembicaraan, serta menentukan tanggal jatuh tempo jadwal tindak lanjut (`follow_up_date`). Sistem secara otomatis menandai status follow-up menjadi `pending` dan menampilkan badge pengingat pada dashboard sales.")

    add_placeholder_box("Gambar 3.7 Timeline Riwayat Aktivitas & Follow-Up Interaksi Penjualan", "[Screenshot Komponen Timeline Riwayat Aktivitas dan Pengingat Follow-up]")

    add_heading_3("3.1.2.7 Penyiaran Pesan Penjualan (Sales Deal Blast Messaging)")
    add_p("Untuk meningkatkan efektivitas penawaran promo dan follow-up massal, penulis mengembangkan fitur penyiaran pesan penjualan (`DealController::blast` & `DealService::blastMessage`). Staf sales dapat memilih beberapa kartu deal aktif di pipeline, memilih saluran komunikasi (WhatsApp atau Email), menuliskan pesan promosi, dan mengunggah gambar brosur treatment. Pesan dikirimkan secara serentak ke nomor WhatsApp atau alamat email prospek yang terhubung, dan setiap penyiaran otomatis dicatat sebagai riwayat aktivitas pada database.")

    add_placeholder_box("Gambar 3.8 Form Siaran Pesan Penjualan (Sales Blast Modal)", "[Screenshot Modal Form Blast WhatsApp & Email Sales]")

    add_heading_3("3.1.2.8 Struktur Routing, Form Request, dan Model Modul Sales")
    add_p("Daftar endpoint rute pada modul sales dikelompokkan dalam middleware `role:Sales` dengan prefix `/sales`, sebagaimana ditunjukkan pada Tabel 3.2.")

    add_custom_table(
        headers=["Method", "URI Rute Modul Sales", "Controller & Method Handler", "Keterangan Fungsi Bisnis"],
        rows=[
            ["GET", "/sales/dashboard", "Sales\\DashboardController@index", "Menampilkan statistik KPI sales, target, dan aktivitas hari ini"],
            ["GET", "/sales/leads", "Sales\\LeadController@index", "Menampilkan daftar seluruh prospek calon pelanggan milik sales"],
            ["POST", "/sales/leads", "Sales\\LeadController@store", "Menyimpan data prospek baru yang diinput sales"],
            ["POST", "/sales/leads/{lead}/qualify", "Sales\\LeadController@qualify", "Memperbarui status kualifikasi prospek (Qualified/Unqualified)"],
            ["POST", "/sales/leads/{lead}/convert", "Sales\\LeadController@convert", "Mengarahkan prospek qualified ke proses pembuatan deal"],
            ["GET", "/sales/pipeline", "Sales\\DealController@pipeline", "Menampilkan papan Kanban visual pipeline transaksi"],
            ["POST", "/sales/deals", "Sales\\DealController@store", "Membuat deal transaksi baru dan mengubah status lead ke converted"],
            ["POST", "/sales/deals/{deal}/move-stage", "Sales\\DealController@moveStage", "Mengubah tahapan pipeline deal hasil drag-and-drop (AJAX)"],
            ["POST", "/sales/deals/{deal}/close", "Sales\\DealController@close", "Menutup deal sebagai Won (buat customer) atau Lost"],
            ["POST", "/sales/deals/blast", "Sales\\DealController@blast", "Mengirimkan pesan penawaran massal via WhatsApp/Email"],
            ["POST", "/sales/activities", "Sales\\ActivityController@store", "Mencatat log interaksi dan jadwal follow-up baru"],
            ["POST", "/sales/activities/{id}/complete", "Sales\\ActivityController@completeFollowUp", "Menandai jadwal follow-up telah selesai dilaksanakan"]
        ],
        caption="Tabel 3.2 Daftar Endpoint Rute Modul Sales"
    )

    add_heading_3("3.1.3 Kendala yang Dihadapi")
    add_p("Selama proses perancangan dan implementasi Modul Sales, I Nyoman Viveka menghadapi beberapa kendala teknis, antara lain:\n"
          "1. Inkonsistensi DOM dan State Sorting pada SortableJS: Ketika kartu transaksi dipindahkan secara cepat antar kolom pipeline, terjadi konflik token nama class pada event handler JavaScript (`InvalidCharacterError`), sehingga badge jumlah transaksi dan akumulasi nominal nilai stage tidak ter-update secara real-time tanpa refresh halaman.\n"
          "2. Penanganan Eksepsi Relasi Notifikasi Deal Closing: Pada rilis awal, event `DealWonNotification` mengalami galat null reference akibat pemanggilan properti relasi usang `$deal->sales` alih-alih menggunakan relasi resmi `$deal->assignedUser`.\n"
          "3. Integritas Transaksi saat Konversi Prospek: Risiko terjadinya inkonsistensi data ketika sistem berhasil membuat entitas `Deal` namun gagal memperbarui status `Lead` karena kegagalan jaringan atau validasi.")

    add_heading_3("3.1.4 Solusi Penyelesaian Kendala")
    add_p("Untuk mengatasi kendala-kendala teknis tersebut, penulis menerapkan solusi-solusi terstruktur sebagai berikut:\n"
          "1. Merombak integrasi SortableJS dengan menggunakan token identifikasi tunggal (`.stage-count-badge`) dan mengembalikan response JSON komprehensif dari backend yang memuat total deal dan nilai terkini dari setiap stage, lalu memperbarui elemen DOM melalui micro-animation Alpine.js.\n"
          "2. Melakukan audit dan standardisasi relasi model Eloquent pada seluruh notifikasi sistem, memastikan `DealWonNotification` merujuk secara konsisten ke `$this->deal->assignedUser`.\n"
          "3. Membungkus seluruh proses konversi prospek dan penutupan transaksi ke dalam blok transaksi atomik `DB::transaction`, sehingga jika terjadi anomali pada salah satu proses, sistem secara otomatis melakukan *rollback* ke kondisi semula.")

    # ═════════════════════════════════════════════════════════════
    # SUBBAB 3.2 ALVIN SAPUTRA
    # ═════════════════════════════════════════════════════════════
    add_heading_2("3.2 Pelaksanaan PKL: ALVIN SAPUTRA, NPM 23312138")

    add_heading_3("3.2.1 Bidang Kerja")
    add_p("Pelaksanaan PKL yang dilakukan oleh Alvin Saputra (NPM 23312138) dilaksanakan di PT Microdata Indonesia pada Divisi Web Development di bawah bimbingan Bapak Sigit Wasis Subekti selaku Programmer. Penulis bertanggung jawab terhadap perancangan Modul Administrator, Arsitektur Keamanan Kontrol Akses (RBAC), Manajemen Data Master, Sistem Impor/Ekspor Massal, Sistem Audit Trail Otomatis, Integrasi Cloud Object Storage MinIO (S3), serta Orkestrasi Kontainerisasi Docker Compose.")

    add_heading_3("3.2.2 Pelaksanaan Kerja")
    add_p("Pelaksanaan kerja yang dilakukan oleh Alvin Saputra dalam membangun fondasi arsitektural dan modul Administrator BEAUTY_CRM diuraikan sebagai berikut:")

    add_heading_3("3.2.2.1 Analisis & Perancangan Kontrol Akses Berbasis Peran (RBAC)")
    add_p("Penulis merancang arsitektur keamanan kontrol akses berbasis peran (*Role-Based Access Control*) menggunakan pustaka `Spatie Laravel Permission 6.x`. Matriks otorisasi sistem membagi hak akses ke dalam 11 permission granular yang dikelompokkan ke dalam empat peran pengguna, sebagaimana dirangkum pada Tabel 3.1.")

    add_custom_table(
        headers=["Hak Akses (Permission)", "Admin", "Sales", "Customer Service", "Manager"],
        rows=[
            ["manage users", "V", "-", "-", "-"],
            ["manage leads", "V", "V", "-", "-"],
            ["manage deals", "V", "V", "-", "-"],
            ["manage customers", "V", "-", "V", "-"],
            ["manage activities", "V", "V", "V", "-"],
            ["manage pipeline", "V", "-", "-", "V"],
            ["manage sources", "V", "-", "-", "-"],
            ["view reports", "V", "-", "-", "V"],
            ["view audit logs", "V", "-", "-", "V"],
            ["import export data", "V", "-", "-", "-"],
            ["manage settings", "V", "-", "-", "-"]
        ],
        caption="Tabel 3.1 Matriks Hak Akses Peran (Role & Permission Matrix) BEAUTY_CRM"
    )

    add_placeholder_box("Gambar 3.15 Diagram Entity Relationship (ERD) Sistem BEAUTY_CRM", "[Bagan ERD 10 Tabel: Users, Leads, Deals, Customers, Activities, AuditLogs, Stages, Sources, Reasons]")

    add_heading_3("3.2.2.2 Manajemen Pengguna & Penetapan Target (User & Target Management)")
    add_p("Penulis membangun modul pengelolaan staf pengguna pada `Admin\\UserController`. Selain operasi CRUD dasar, penulis menambahkan fitur krusial berupa penetapan target kerja bulanan bagi setiap staf:")
    add_p("1. `monthly_target`: Target kuantitatif jumlah perolehan member baru / customer baru per bulan.\n"
          "2. `revenue_target`: Target nominal perolehan omset penjualan per bulan dalam mata uang Rupiah.")

    add_p("Penulis juga menambahkan fitur *Toggle Active/Inactive* via AJAX untuk menonaktifkan akun staf yang telah resign tanpa menghapus data historis transaksi mereka. Sistem dilengkapi proteksi keamanan *self-delete prevention* agar akun Administrator utama tidak dapat dinonaktifkan atau dihapus oleh dirinya sendiri.")

    add_placeholder_box("Gambar 3.9 Antarmuka Manajemen Akun Pengguna & Pengaturan Target Performa", "[Screenshot Halaman Admin User Management dengan Tabel Staf dan Modal Target]")

    add_heading_3("3.2.2.3 Pengelolaan Data Master & Reordering Tahapan Pipeline")
    add_p("Penulis mengembangkan panel pengelolaan data master pendukung CRM:")
    add_p("1. **Lead Sources**: Mengelola saluran perolehan prospek (WhatsApp Ads, Instagram Campaign, Website, Referral, Event Kecantikan, Walk-in) dengan konfigurasi warna badge dan ikon visual (`LeadSourceController`).\n"
          "2. **Lost Reasons**: Mengelola kamus alasan kegagalan transaksi guna keperluan evaluasi analitik manajerial (`LostReasonController`).\n"
          "3. **Pipeline Stages**: Mengelola tahapan pipeline penjualan beserta bobot persentase probabilitas closing (0-100%). Penulis mengimplementasikan fitur **Drag & Drop Reordering** urutan tahapan menggunakan SortableJS yang mengirimkan urutan indeks baru ke endpoint `PATCH /admin/pipeline-stages/reorder`.")

    add_placeholder_box("Gambar 3.10 Antarmuka Penataan Ulang Tahapan Pipeline (Drag & Drop Reorder)", "[Screenshot Pengaturan Tahapan Pipeline dengan Fitur Drag Reorder Urutan Stage]")

    add_heading_3("3.2.2.4 Impor dan Ekspor Data Massal Berbasis Excel (Maatwebsite Excel)")
    add_p("Untuk mempermudah migrasi data calon pelanggan dalam jumlah ribuan dari klinik mitra, penulis mengimplementasikan sistem impor data prospek massal menggunakan pustaka `Maatwebsite Excel 3.x` pada class `LeadImport` dan `ImportExportService`.")

    add_p("Fitur impor dilengkapi validasi *Chunk Reading* (500 baris per batch), *Batch Insert* (100 record per query), pencocokan otomatis nama sumber prospek (*source name matching*), serta penanganan khusus untuk melakukan casting otomatis pada nomor telepon seluler bertipe data numerik Excel agar tidak memicu kegagalan validasi string.")

    add_placeholder_box("Gambar 3.11 Antarmuka Fitur Impor Data Prospek Massal via Excel", "[Screenshot Halaman Import Leads dengan Upload Form dan Download Template Button]")

    add_heading_3("3.2.2.5 Sistem Rekam Jejak Audit Otomatis (Polymorphic Audit Observer)")
    add_p("Guna memenuhi standar tata kelola kepatuhan data medis dan finansial klinik, penulis merancang sistem rekam jejak audit (*Audit Trail*) otomatis menggunakan pola *Observer Pattern* pada `app/Observers/AuditObserver.php`. Observer ini didaftarkan melalui `app/Providers/AuditServiceProvider.php` untuk memantau siklus hidup empat model utama: `User`, `Lead`, `Deal`, dan `Customer`.")

    add_p("Setiap kali terjadi aksi `created`, `updated`, atau `deleted`, Observer secara otomatis mengekstraksi nilai data sebelum (`old_values`) dan sesudah perubahan (`new_values`), menyaring field sensitif (seperti password dan token sesi), serta merekam identitas pengguna yang melakukan mutasi, alamat IP client, dan User Agent browser:")

    add_p("```php\n"
          "private function log(string $action, Model $model, ?array $old, ?array $new): void {\n"
          "    AuditLog::create([\n"
          "        'user_id'        => Auth::id(),\n"
          "        'action'         => $action,\n"
          "        'auditable_type' => get_class($model),\n"
          "        'auditable_id'   => $model->getKey(),\n"
          "        'old_values'     => $old ? $this->filterFields($old) : null,\n"
          "        'new_values'     => $new ? $this->filterFields($new) : null,\n"
          "        'ip_address'     => Request::ip(),\n"
          "        'user_agent'     => Request::userAgent(),\n"
          "    ]);\n"
          "}\n"
          "```", italic=True)

    add_placeholder_box("Gambar 3.12 Tampilan Tabel Riwayat Rekam Jejak Audit Sistem (Audit Trail)", "[Screenshot Halaman Audit Logs dengan Paginasi dan Detail Perubahan JSON Diff]")

    add_heading_3("3.2.2.6 Integrasi Cloud Object Storage MinIO (S3 Compatible Storage)")
    add_p("Penulis mengintegrasikan MinIO sebagai media penyimpanan objek berbasis protokol Amazon S3 untuk menangani unggahan berkas foto avatar profil pengguna dan lampiran siaran pesan promosi. Penulis mengonfigurasi driver filesystem S3 pada `config/filesystems.php` dan mendefinisikan variabel environment pada `.env.docker`:")

    add_custom_table(
        headers=["Variabel Environment", "Nilai Konfigurasi", "Fungsi dan Keterangan"],
        rows=[
            ["FILESYSTEM_DISK", "s3", "Menetapkan driver default penyimpanan file Laravel ke protokol S3"],
            ["AWS_ACCESS_KEY_ID", "minioadmin", "Kunci akses otentikasi akun root MinIO"],
            ["AWS_SECRET_ACCESS_KEY", "minioadmin", "Kunci rahasia otentikasi akun root MinIO"],
            ["AWS_DEFAULT_REGION", "us-east-1", "Wilayah region default penyimpanan S3"],
            ["AWS_BUCKET", "beauty-crm", "Nama bucket utama penyimpanan objek aplikasi"],
            ["AWS_ENDPOINT", "http://minio:9000", "URL endpoint internal komunikasi antar-kontainer Docker"],
            ["AWS_URL", "http://localhost:9010/beauty-crm", "URL endpoint publik yang diakses oleh browser pengguna"],
            ["AWS_USE_PATH_STYLE_ENDPOINT", "true", "Mengaktifkan format akses berbasis path style endpoint"]
        ],
        caption="Tabel 3.7 Parameter Konfigurasi Environment Object Storage MinIO S3"
    )

    add_placeholder_box("Gambar 3.13 Dashboard Manajemen Penyimpanan Objek MinIO Console", "[Screenshot Dashboard Web MinIO Console Port 9011 dengan Bucket beauty-crm]")

    add_heading_3("3.2.2.7 Orkestrasi Kontainerisasi Sistem dengan Docker Compose")
    add_p("Untuk menyederhanakan proses deployment dan menjamin keseragaman lingkungan antara tahap pengembangan (*development*) dan produksi (*production*), penulis merancang konfigurasi kontainerisasi multi-service menggunakan Docker Compose pada berkas `docker-compose.yml`.")

    add_custom_table(
        headers=["Nama Kontainer Layanan", "Base Image", "Port Binding", "Peran dan Tanggung Jawab Layanan"],
        rows=[
            ["beauty_crm_app", "Custom Build (PHP 8.2-FPM)", "Internal (9000)", "Menjalankan runtime aplikasi Laravel 12 beserta seluruh ekstensi PHP yang dibutuhkan"],
            ["beauty_crm_nginx", "nginx:1.25-alpine", "8080:80", "Web server reverse proxy yang menerima request HTTP dari browser dan meneruskan ke PHP-FPM"],
            ["beauty_crm_minio", "minio/minio:latest", "9010:9000\n9011:9001", "Penyedia layanan Object Storage kompatibel S3 (Port 9010 API, Port 9011 Console)"],
            ["beauty_crm_minio_setup", "minio/mc:latest", "Internal Service", "Job otomatis yang menginisialisasi pembuatan bucket beauty-crm dan menetapkan kebijakan public"]
        ],
        caption="Tabel 3.6 Konfigurasi Port & Layanan Docker Compose"
    )

    add_placeholder_box("Gambar 3.14 Diagram Arsitektur Kontainerisasi Layanan Docker Compose", "[Bagan Kontainer: Nginx -> App PHP-FPM -> MinIO S3 & Host MySQL]")

    add_p("Penulis juga menyusun berkas script instalasi otomatis `docker/setup.sh` yang mengeksekusi penyalinan environment, build kontainer, migrasi database dengan flag `--force`, pembuatan symlink storage, dan pembersihan cache sistem dalam satu perintah eksekusi tunggal.")

    add_heading_3("3.2.2.8 Struktur Routing, Form Request, dan Model Modul Administrator")
    add_p("Rute-rute administrasi sistem diproteksi di bawah middleware `role:Admin` pada berkas `routes/admin.php`, sebagaimana dirangkum pada Tabel 3.3.")

    add_custom_table(
        headers=["Method", "URI Rute Modul Admin", "Controller & Handler", "Fungsi Operasional Sistem"],
        rows=[
            ["GET", "/admin/dashboard", "Admin\\DashboardController@index", "Menampilkan statistik ringkasan global seluruh entitas CRM"],
            ["GET|POST", "/admin/users", "Admin\\UserController@index|store", "Manajemen CRUD data akun pengguna dan target bulanan"],
            ["PATCH", "/admin/users/{user}/toggle", "Admin\\UserController@toggle", "Mengubah status aktif/nonaktif akun staf secara instan"],
            ["GET|POST", "/admin/lead-sources", "Admin\\LeadSourceController@index|store", "Pengelolaan master data sumber prospek calon pelanggan"],
            ["PATCH", "/admin/pipeline-stages/reorder", "Admin\\PipelineStageController@reorder", "Menyimpan urutan baru tahapan pipeline hasil drag reorder"],
            ["GET|POST", "/admin/pipeline-stages", "Admin\\PipelineStageController@index|store", "Pengelolaan master tahapan pipeline dan probabilitas"],
            ["GET|POST", "/admin/lost-reasons", "Admin\\LostReasonController@index|store", "Pengelolaan master kamus alasan kegagalan transaksi"],
            ["POST", "/admin/leads/import", "Admin\\ImportExportController@import", "Mengeksekusi impor data prospek massal dari file Excel"],
            ["GET", "/admin/leads/export", "Admin\\ImportExportController@export", "Mengunduh data prospek terdaftar ke dalam format Excel"],
            ["GET", "/admin/audit-logs", "Admin\\AuditLogController@index", "Menampilkan visualisasi rekam jejak audit mutasi data"],
            ["GET|PUT", "/admin/settings", "Admin\\SettingsController@index|update", "Konfigurasi profil instansi dan toggle notifikasi sistem"]
        ],
        caption="Tabel 3.3 Daftar Endpoint Rute Modul Administrator"
    )

    add_heading_3("3.2.3 Kendala yang Dihadapi")
    add_p("Dalam merancang infrastruktur dan modul Admin, Alvin Saputra menghadapi kendala-kendala berikut:\n"
          "1. Masalah Komunikasi Jaringan Docker ke Database Host Lokal: Kontainer PHP-FPM di dalam lingkungan Docker network tidak dapat langsung menghubungi database MySQL yang berjalan pada server Laragon lokal host machine via IP `127.0.0.1`.\n"
          "2. Galat Tipe Data Sel Numerik pada Impor Excel: Ketika pengguna mengimpor file Excel yang berisi kolom nomor telepon berformat angka besar (misal `6281234567890`), library PhpSpreadsheet membaca nilai tersebut sebagai *float/integer scientific notation*, yang menyebabkan kegagalan validasi string pada Laravel FormRequest.\n"
          "3. Perbedaan URL Penyimpanan S3 antara Lingkungan Internal Kontainer dan Browser: Hostname internal MinIO (`http://minio:9000`) yang digunakan oleh PHP-FPM tidak dapat diakses langsung oleh browser client publik yang mengakses dari `http://localhost:9010`.")

    add_heading_3("3.2.4 Solusi Penyelesaian Kendala")
    add_p("Solusi teknis yang diterapkan oleh penulis untuk mengatasi kendala di atas adalah:\n"
          "1. Menambahkan pemetaan host khusus `host.docker.internal:host-gateway` pada service `app` di `docker-compose.yml`, sehingga kontainer PHP dapat mengakses port 3306 MySQL pada sistem operasi host secara mulus.\n"
          "2. Mengimplementasikan method transformasi tipe data pada `LeadImport` dengan melakukan *string casting* dan normalisasi format nomor telepon seluler sebelum dialirkan ke tahap validasi Laravel.\n"
          "3. Memisahkan konfigurasi `AWS_ENDPOINT` (untuk komunikasi backend internal via Docker DNS) dengan `AWS_URL` (untuk pembentukan URL publik aset media yang diakses oleh browser pengguna pada port 9010).")

    # ═════════════════════════════════════════════════════════════
    # SUBBAB 3.3 DIMAS APRIANTO
    # ═════════════════════════════════════════════════════════════
    add_heading_2("3.3 Pelaksanaan PKL: DIMAS APRIANTO, NPM 23312136")

    add_heading_3("3.3.1 Bidang Kerja")
    add_p("Pelaksanaan PKL yang dilakukan oleh Dimas Aprianto (NPM 23312136) dilaksanakan di PT Microdata Indonesia pada Divisi Web Development di bawah bimbingan Bapak Sigit Wasis Subekti selaku Programmer. Penulis berfokus pada pengembangan Modul Eksekutif Manager (BI Reports & Forecasting) serta Modul Layanan Pelanggan (Customer Service & Multi-Channel Broadcast). Penulis bertanggung jawab atas perancangan arsitektur Service Layer pelaporan bisnis, algoritma peramalan pendapatan berbasis probabilitas tahapan transaksi, manajemen data pelanggan purna jual, pemantauan jadwal follow-up, serta integrasi gateway WhatsApp Fonnte dan email broadcast.")

    add_heading_3("3.3.2 Pelaksanaan Kerja")
    add_p("Pelaksanaan kerja yang dilakukan oleh Dimas Aprianto diuraikan ke dalam beberapa sub-komponen teknis sebagai berikut:")

    add_heading_3("3.3.2.1 Analisis Business Intelligence & Customer Service Workflow")
    add_p("Penulis menganalisis kebutuhan eksekutif manajemen klinik kecantikan yang memerlukan laporan analitik performa komprehensif tanpa membebani performa database transaksional. Penulis juga menganalisis alur kerja divisi Customer Service yang bertugas menjaga hubungan jangka panjang dengan pelanggan yang telah melakukan perawatan (*post-treatment customer retention*).")

    add_placeholder_box("Gambar 3.16 Tampilan Dashboard Utama Eksekutif Manager (BI Dashboard)", "[Screenshot Dashboard Manager dengan Visualisasi KPI, Chart Funnel, dan Leaderboard]")

    add_heading_3("3.3.2.2 Perancangan Arsitektur Report Center (ReportService)")
    add_p("Penulis merancang class arsitektural `app/Services/ReportService.php` yang memuat lebih dari 800 baris logika bisnis murni. Pendekatan Service Layer ini memisahkan perhitungan matematika, statistik, peramalan, dan aggregasi data dari controller:")
    add_p("1. `getManagerDashboard()`: Mengagregasi metrik KPI utama (Total Deals, Total Pendapatan, Top Treatment Product, Rasio Pencapaian Target, Funnel Konversi, dan Leaderboard).\n"
          "2. `getSalesPerformance()`: Menghitung total prospek, total transaksi, total deal won, win rate persentase, dan rata-rata durasi penutupan transaksi (*average days to close*) per tenaga penjual.\n"
          "3. `getRevenueReport()`: Menghitung tren pendapatan historis selama 12 bulan terakhir dengan filter rentang waktu dinamis.\n"
          "4. `getLostReasons()`: Menganalisis penyebab utama kegagalan penutupan transaksi dalam bentuk persentase kontribusi.\n"
          "5. `getLeadSources()`: Mengukur rasio efektivitas dan pengembalian investasi (ROI) dari setiap saluran perolehan prospek.")

    add_heading_3("3.3.2.3 Metrik Performa Sales, Rasio Konversi, dan Tren Pendapatan")
    add_p("Penulis mengimplementasikan perhitungan rasio kemenangan (*Win Rate*) dan durasi rata-rata penutupan transaksi (*Average Days to Close*) dengan rumus:")
    add_p("$$\\text{Win Rate (\\%)} = \\left( \\frac{\\text{Total Deal Won}}{\\text{Total Deal Won} + \\text{Total Deal Lost}} \\right) \\times 100\\%$$", italic=True)
    add_p("$$\\text{Average Days to Close} = \\frac{\\sum (\\text{closed\\_at} - \\text{created\\_at})}{\\text{Total Deal Won}}$$", italic=True)

    add_placeholder_box("Gambar 3.17 Laporan Analisis Performa Tim Sales & Conversion Win Rate", "[Screenshot Halaman Laporan Sales Performance dengan Tabel Metrik dan Win Rate]")
    add_placeholder_box("Gambar 3.18 Diagram Analisis Tren Pendapatan 12 Bulan Historis", "[Screenshot Grafik Batang Tren Pendapatan 12 Bulan Terakhir]")

    add_heading_3("3.3.2.4 Mesin Peramalan Pendapatan & Member (Weighted Pipeline Forecast)")
    add_p("Salah satu kontribusi inovatif dari penulis adalah perancangan Mesin Peramalan Penjualan (*Forecast Engine*) pada `ForecastController` dan `ReportService::getForecastData()`. Mesin ini memproyeksikan potensi pendapatan klinik di masa mendatang berdasarkan nilai transaksi terbobot (*Weighted Pipeline Value*).")

    add_p("Setiap transaksi aktif yang berada di tahapan pipeline dikalikan dengan persentase probabilitas keberhasilan dari tahapan tersebut, sebagaimana dirumuskan pada Tabel 3.8.")

    add_custom_table(
        headers=["Tahapan Pipeline (Stage Name)", "Probabilitas Default", "Rumus Perhitungan Nilai Terbobot (Weighted Value)"],
        rows=[
            ["Prospecting / Initial Contact", "10 %", "Nilai Rata-rata Transaksi x 0.10"],
            ["Beauty Consultation / Qualification", "30 %", "Nilai Rata-rata Transaksi x 0.30"],
            ["Treatment Proposal Sent", "60 %", "Nilai Rata-rata Transaksi x 0.60"],
            ["Price & Schedule Negotiation", "80 %", "Nilai Rata-rata Transaksi x 0.80"],
            ["Final Closing (Won)", "100 %", "Nilai Transaksi Aktual x 1.00"]
        ],
        caption="Tabel 3.8 Parameter Probabilitas Tahapan Pipeline Deals"
    )

    add_placeholder_box("Gambar 3.19 Tampilan Antarmuka Mesin Peramalan (Weighted Forecasting Board)", "[Screenshot Halaman Sales Forecast dengan Matriks Weighted Value dan Proyeksi]")

    add_heading_3("3.3.2.5 Basis Data Pelanggan, Profiling, dan Segmentasi Nilai Belanja")
    add_p("Pada Modul Customer Service (`CS\\CustomerController` dan `CustomerService`), penulis mengelola basis data pelanggan yang telah terkonversi dari deal won. Sistem dilengkapi dengan kalkulasi akumulasi total belanja (*total spend*) pelanggan secara dinamis melalui subquery SQL:")

    add_p("```php\n"
          "public function getCustomers(array $filters = []): LengthAwarePaginator {\n"
          "    return Customer::with('csUser')\n"
          "        ->select('customers.*')\n"
          "        ->selectRaw(\"IFNULL((SELECT SUM(value) FROM deals WHERE deals.lead_id = customers.lead_id AND deals.status = 'won'), 0) as total_spend\")\n"
          "        ->selectRaw(\"(SELECT MAX(created_at) FROM activities WHERE activities.activitable_id = customers.id AND activities.activitable_type = ?) as last_contacted_at\", [Customer::class])\n"
          "        ->search($filters['search'] ?? null)\n"
          "        ->when($filters['min_spend'] ?? null, fn($q, $v) => $q->minSpend($v))\n"
          "        ->latest()\n"
          "        ->paginate(15);\n"
          "}\n"
          "```", italic=True)

    add_p("Pelanggan dapat dikelompokkan dengan penandaan tag JSON fleksibel (misal: *VIP Member, Treatment Acne, Loyalty Silver, Sensitive Skin*) untuk mempermudah strategi penawaran perawatan yang dipersonalisasi.")

    add_placeholder_box("Gambar 3.20 Antarmuka Manajemen Data Pelanggan & Segmentasi Belanja", "[Screenshot Halaman CS Customer List dengan Badge Tags dan Kolom Total Spend]")

    add_heading_3("3.3.2.6 Sistem Manajemen Jadwal Tindak Lanjut & Monitoring Overdue")
    add_p("Penulis membangun dashboard pemantauan follow-up operasional pada `CS\\FollowUpController`. Sistem memisahkan jadwal tindak lanjut ke dalam tiga tab klasifikasi waktu:")
    add_p("1. **Follow-Up Hari Ini**: Daftar pelanggan yang harus dihubungi pada tanggal berjalan.\n"
          "2. **Follow-Up Overdue**: Pengingat kritis untuk jadwal tindak lanjut yang telah melewati tanggal jatuh tempo namun belum diselesaikan oleh staf CS.\n"
          "3. **Follow-Up Selesai**: Arsip riwayat follow-up yang telah tuntas dilaksanakan beserta catatan hasil pembicaraan.")

    add_placeholder_box("Gambar 3.21 Dashboard Jadwal Tindak Lanjut & Monitoring Keterlambatan CS", "[Screenshot Halaman CS Follow-ups dengan Tab Pending, Overdue, dan Completed]")

    add_heading_3("3.3.2.7 Integrasi Multi-Channel Blast (WhatsApp Fonnte API & Email CID)")
    add_p("Untuk mendukung kampanye perawatan berkala dan promosi klinik, penulis merancang sistem penyiaran pesan massal multi-saluran pada `CustomerService::blastMessage()`:")
    add_p("1. **WhatsApp Gateway (Fonnte API)**: Penulis membangun parser otomatis yang mengonversi format HTML dari form web menjadi format WhatsApp Markdown (*bold*, _italic_, ~strike~). Penulis mengintegrasikan cURL transport resmi Fonnte dengan class `CURLFile` untuk mentransmisikan berkas gambar promosi secara langsung ke server WhatsApp Gateway.\n"
          "2. **Email Broadcast (Content-ID Inline Image)**: Penulis membangun mailable class `app/Mail/BlastMessageMail.php`. Untuk memastikan gambar promosi tampil sempurna di aplikasi email client penerima tanpa memicu peringatan blokir eksternal, gambar di-embed secara langsung ke dalam badan email menggunakan mekanisme *CID (Content-ID) inline attachment*.\n"
          "3. Setiap aksi penyiaran pesan secara otomatis diinjeksi ke dalam tabel `activities` sebagai bukti rekam interaksi resmi antara klinik dan pelanggan.")

    add_placeholder_box("Gambar 3.22 Antarmuka Siaran Pesan Multi-Saluran CS (WhatsApp & Email Blast)", "[Screenshot Modal CS Broadcast WhatsApp & Email dengan Opsi Gambar Lampiran]")

    add_heading_3("3.3.2.8 Struktur Routing, Form Request, dan Model Modul Manager & CS")
    add_p("Daftar rute modul Manager dan Customer Service diproteksi oleh middleware peran masing-masing, sebagaimana disajikan pada Tabel 3.4.")

    add_custom_table(
        headers=["Method", "URI Rute Modul Manager & CS", "Controller & Handler", "Fungsi Operasional Sistem"],
        rows=[
            ["GET", "/manager/dashboard", "Manager\\DashboardController@index", "Dashboard analitik utama eksekutif dengan KPI dan grafik tren"],
            ["GET", "/manager/reports/sales-performance", "Manager\\ReportController@salesPerformance", "Laporan performa individual tim sales dan rasio win rate"],
            ["GET", "/manager/reports/revenue", "Manager\\ReportController@revenue", "Laporan analisis tren pendapatan 12 bulan historis"],
            ["GET", "/manager/reports/lost-reasons", "Manager\\ReportController@lostReasons", "Laporan distribusi alasan kegagalan transaksi penjualan"],
            ["GET", "/manager/reports/lead-sources", "Manager\\ReportController@leadSources", "Laporan evaluasi efektivitas saluran perolehan prospek"],
            ["GET", "/manager/reports/export", "Manager\\ReportController@export", "Mengekspor laporan performa sales/revenue ke file Excel"],
            ["GET", "/manager/forecast", "Manager\\ForecastController@index", "Menampilkan visualisasi peramalan nilai terbobot pipeline"],
            ["GET", "/cs/dashboard", "CS\\DashboardController@index", "Dashboard operasional CS dengan ringkasan tiket dan follow-up"],
            ["GET|POST", "/cs/customers", "CS\\CustomerController@index|store", "Manajemen basis data pelanggan dan pencatatan riwayat"],
            ["POST", "/cs/customers/blast", "CS\\CustomerController@blast", "Mengeksekusi siaran pesan massal via WhatsApp dan Email"],
            ["GET", "/cs/follow-ups", "CS\\FollowUpController@index", "Menampilkan jadwal tindak lanjut harian, overdue, dan selesai"],
            ["POST", "/cs/follow-ups/{id}/complete", "CS\\FollowUpController@complete", "Memverifikasi penyelesaian jadwal follow-up pelanggan"]
        ],
        caption="Tabel 3.4 Daftar Endpoint Rute Modul Manager & Customer Service"
    )

    add_heading_3("3.3.3 Kendala yang Dihadapi")
    add_p("Dalam mengembangkan Modul Manager dan Customer Service, Dimas Aprianto menghadapi kendala-kendala berikut:\n"
          "1. Masalah Kinerja N+1 Query pada Agregasi Laporan: Pada rilis awal, method perhitungan performa sales mengeksekusi lebih dari 6 subquery per salesperson di dalam loop `map()`, yang menyebabkan degradasi performa ketika volume data membesar.\n"
          "2. Format Rendering WhatsApp Markdown vs Rich Text HTML: Teks pesan promosi yang diinputkan dari antarmuka web memuat tag HTML (`<p>`, `<strong>`, `<br>`), yang apabila dikirim langsung ke WhatsApp akan tampil sebagai tag mentah yang mengganggu kenyamanan baca pelanggan.\n"
          "3. Pemblokiran Gambar Eksternal pada Email Client: Pengiriman email yang menyertakan link gambar eksternal sering kali diblokir oleh aplikasi email client (seperti Gmail dan Outlook) demi alasan keamanan privasi penerima.")

    add_heading_3("3.3.4 Solusi Penyelesaian Kendala")
    add_p("Solusi teknis yang diimplementasikan oleh penulis adalah:\n"
          "1. Mengoptimalkan query database menggunakan eager loading berelasi (`withCount`, `withSum`) dan subquery SQL murni (`selectRaw`), mereduksi jumlah query dari puluhan menjadi satu kali query tunggal yang efisien.\n"
          "2. Membangun method parser regex `formatMessageForWhatsApp()` pada Service Layer yang secara cerdas mengubah tag `<strong>` menjadi `*bold*`, `<em>` menjadi `_italic_`, `<del>` menjadi `~strike~`, dan `<br>` menjadi baris baru.\n"
          "3. Menggunakan mekanisme Content-ID (CID) inline embedding pada `BlastMessageMail`, di mana berkas gambar diunggah dan disematkan langsung sebagai *MIME multipart part* di dalam paket email, sehingga gambar dapat langsung ditampilkan tanpa memerlukan URL eksternal.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # BAB IV PENUTUP
    # ─────────────────────────────────────────────────────────────
    add_heading_1("BAB IV\nPENUTUP")

    add_heading_2("4.1 Kesimpulan")
    add_p("Berdasarkan hasil pelaksanaan PKL di PT Microdata Indonesia yang dilaksanakan pada tanggal 6 Juli 2026 sampai dengan 6 September 2026, maka dapat disimpulkan bahwa tujuan PKL telah tercapai dengan baik. Adapun kesimpulan yang diperoleh adalah sebagai berikut:")
    add_p("1. Selama pelaksanaan PKL, penulis berhasil menerapkan berbagai ilmu yang diperoleh selama perkuliahan, khususnya pada bidang pemrograman web, basis data, arsitektur perangkat lunak, dan kontrol akses berbasis peran. Implementasi tersebut dilakukan melalui pengembangan sistem BEAUTY_CRM menggunakan Laravel 12, Tailwind CSS 4, Alpine.js, Spatie Permission, Docker, dan MinIO S3.\n"
          "2. Melalui keterlibatan langsung dalam proyek BEAUTY_CRM, penulis memperoleh pemahaman mengenai alur kerja pengembangan perangkat lunak secara profesional. Proses tersebut meliputi analisis kebutuhan sistem CRM, perancangan database relasional, pengembangan logika bisnis Service Layer, integrasi layanan pihak ketiga (WhatsApp Fonnte API & S3 Object Storage), pengujian sistem, hingga orkestrasi deployment menggunakan Docker Compose. Selain itu, penulis juga memahami pola kerja tim dalam lingkungan industri teknologi informasi yang menerapkan pembagian modul dan koordinasi antar anggota tim pengembang.\n"
          "3. Selama PKL, penulis memperoleh pengalaman dalam menyelesaikan berbagai permasalahan teknis, seperti optimasi performa query database, penanganan data numerik Excel, sinkronisasi state interaktif Kanban SortableJS, hingga pengelolaan penyimpanan cloud MinIO. Selain meningkatkan kemampuan teknis, kegiatan PKL juga membantu penulis mengembangkan kemampuan komunikasi, kerja sama tim, manajemen waktu, tanggung jawab, serta kemampuan beradaptasi terhadap lingkungan kerja profesional.")

    add_heading_2("4.2 Saran")
    add_p("Berdasarkan hasil pelaksanaan PKL di PT Microdata Indonesia, terdapat beberapa saran yang dapat diberikan sebagai berikut, yaitu:")
    add_p("1. Universitas Teknokrat Indonesia\n"
          "   Universitas Teknokrat Indonesia diharapkan dapat terus memperluas kerja sama dengan perusahaan berbasis teknologi informasi seperti PT Microdata Indonesia agar mahasiswa memiliki lebih banyak pilihan tempat PKL yang relevan dengan bidang studinya, serta terus memperkaya materi praktikum perkuliahan dengan teknologi modern seperti Docker dan Cloud Storage.\n"
          "2. PT Microdata Indonesia\n"
          "   PT Microdata Indonesia diharapkan dapat terus memberikan kesempatan kepada mahasiswa untuk terlibat dalam proyek nyata. Hal ini bertujuan agar mahasiswa mampu mengasah keterampilan analisis, problem-solving, serta beradaptasi dengan kebutuhan industri yang dinamis.\n"
          "3. Rekomendasi bagi Mahasiswa\n"
          "   Mahasiswa yang akan melaksanakan PKL disarankan untuk mempersiapkan pemahaman fundamental pemrograman, penguasaan framework, basis data relasional, dan penggunaan Git Version Control System secara matang sebelum terjun ke dunia industri.")

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # DAFTAR PUSTAKA
    # ─────────────────────────────────────────────────────────────
    add_heading_1("DAFTAR PUSTAKA")

    pustaka_items = [
        "Ariawan, M. D., Triayudi, A., & Sholihati, I. D. (2020). Perancangan User Interface dan User Experience Aplikasi E-Ticketing Berbasis Web Menggunakan Metode Design Thinking. Jurnal Riset Informatika, 2(4), 163-170.",
        "Bass, L., Clements, P., & Kazman, R. (2021). Software Architecture in Practice (4th ed.). Boston: Addison-Wesley Professional.",
        "Chacon, S., & Straub, B. (2020). Pro Git: Everything You Need to Know About Git (2nd ed.). New York: Apress.",
        "Docker Inc. (2026). Docker Documentation: Containerize Applications with Docker Compose. Diakses pada 15 Agustus 2026, dari https://docs.docker.com/compose/.",
        "Fonnte Technologies. (2026). Fonnte WhatsApp API Gateway Documentation and Integration Guide. Diakses pada 20 Agustus 2026, dari https://fonnte.com/docs/.",
        "Fowler, M. (2018). Patterns of Enterprise Application Architecture. Boston: Addison-Wesley Professional.",
        "Laravel LLC. (2026). Laravel 12.x Documentation: The PHP Framework for Web Artisans. Diakses pada 10 Agustus 2026, dari https://laravel.com/docs/12.x/.",
        "Maatwebsite. (2026). Laravel Excel Documentation (v3.1): Supercharged Excel exports and imports in Laravel. Diakses pada 12 Agustus 2026, dari https://docs.laravel-excel.com/.",
        "Mal, L. H., & Mertayasa, I. N. E. (2018). Pengaruh Diskon dan Kualitas Pelayanan Terhadap Keputusan Pembelian Tiket Pesawat Secara Online. Jurnal Manajemen dan Bisnis, 15(2), 112-125.",
        "MinIO Inc. (2026). MinIO High Performance S3 Compatible Object Storage Documentation. Diakses pada 18 Agustus 2026, dari https://min.io/docs/.",
        "Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). New York: McGraw-Hill Education.",
        "Purbo, Y. S., Utomo, F. S., & Purwati, N. (2023). Pengaruh Praktik Kerja Lapangan Terhadap Kesiapan Kerja Mahasiswa. Jurnal Pendidikan Teknologi dan Kejuruan, 19(1), 45-54.",
        "Rafiqh, M., & Ismail, I. (2023). Analisis User Experience Pada Aplikasi E-Commerce Menggunakan Metode Usability Testing. Jurnal Sistem Informasi dan Informatika, 5(2), 78-86.",
        "Spatie. (2026). Spatie Laravel Permission Package Documentation (v6.x). Diakses pada 8 Agustus 2026, dari https://spatie.be/docs/laravel-permission/.",
        "Universitas Teknokrat Indonesia. (2025). Buku Panduan Pelaksanaan dan Penyusunan Laporan Praktik Kerja Lapangan (PKL) Fakultas Teknik dan Ilmu Komputer. Bandar Lampung: Universitas Teknokrat Indonesia."
    ]

    for p_item in pustaka_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(p_item)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────
    # LAMPIRAN
    # ─────────────────────────────────────────────────────────────
    add_heading_1("LAMPIRAN")

    add_heading_2("Lampiran 1: Formulir Penilaian PKL Mahasiswa")
    add_placeholder_box("Lampiran 1.1 Formulir Penilaian PKL: I Nyoman Viveka (NPM: 23312003)", "[Formulir Penilaian Pembimbing Lapangan - Nilai & Tanda Tangan]")
    add_placeholder_box("Lampiran 1.2 Formulir Penilaian PKL: Alvin Saputra (NPM: 23312138)", "[Formulir Penilaian Pembimbing Lapangan - Nilai & Tanda Tangan]")
    add_placeholder_box("Lampiran 1.3 Formulir Penilaian PKL: Dimas Aprianto (NPM: 23312136)", "[Formulir Penilaian Pembimbing Lapangan - Nilai & Tanda Tangan]")

    add_heading_2("Lampiran 2: Catatan Harian (Logbook) Kegiatan PKL")
    add_placeholder_box("Lampiran 2.1 Catatan Harian Logbook: I Nyoman Viveka", "[Rekap Kegiatan Mingguan 6 Juli 2026 - 6 September 2026]")
    add_placeholder_box("Lampiran 2.2 Catatan Harian Logbook: Alvin Saputra", "[Rekap Kegiatan Mingguan 6 Juli 2026 - 6 September 2026]")
    add_placeholder_box("Lampiran 2.3 Catatan Harian Logbook: Dimas Aprianto", "[Rekap Kegiatan Mingguan 6 Juli 2026 - 6 September 2026]")

    add_heading_2("Lampiran 3: Foto Dokumentasi Kegiatan PKL di PT Microdata Indonesia")
    add_placeholder_box("Lampiran 3.1 Foto Kegiatan Briefing dan Diskusi Arsitektur Sistem Bersama Pembimbing Lapangan", "[Foto Tim Mahasiswa dan Pembimbing Lapangan di Ruang Kerja PT Microdata Indonesia]")
    add_placeholder_box("Lampiran 3.2 Foto Suasana Pengerjaan Proyek BEAUTY_CRM di Kantor PT Microdata Indonesia", "[Foto Mahasiswa Mengerjakan Modul di Meja Kerja]")

    add_heading_2("Lampiran 4: Surat Tugas & Keterangan Selesai PKL dari PT Microdata Indonesia")
    add_placeholder_box("Lampiran 4.1 Surat Keterangan Selesai PKL Resmi dari PT Microdata Indonesia", "[Surat Keterangan Selesai Magang Bertanda Tangan Direktur & Cap Perusahaan]")

    # Save docx
    docx_path = "docs/Laporan_PKL_BEAUTY_CRM.docx"
    doc.save(docx_path)
    print(f"Successfully generated Word document: {docx_path}")

if __name__ == "__main__":
    generate_report()
